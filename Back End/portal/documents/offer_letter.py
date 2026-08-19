"""
OFFER LETTER GENERATION
==============================================================================

Produces the two files HR-OPS receives once HR-APP has signed:

    build_offer_letter_pdf(context)   the official letter, signed, for printing
    build_offer_letter_docx(context)  an editable copy for correcting mistakes

Both return raw bytes. Neither touches the database, the filesystem or Django
settings, and neither imports a model. Everything they need arrives in one
plain dictionary. That is what lets the layout be checked by running
preview_offer_letter.py without a database, a server or an application.

------------------------------------------------------------------------------
PLAIN PAPER
------------------------------------------------------------------------------
The letter is drawn IN FULL, letterhead included, and prints on ordinary A4.
No pre-printed stationery is needed, and the Word version carries the same
header so a corrected letter prints identically.

The header's Hindi lines are pre-rendered images rather than text. That is not
a shortcut -- see the long note on _draw_letterhead() for why drawing them as
text produces visibly wrong Devanagari.

------------------------------------------------------------------------------
WHY THE WORD FILE IS UNSIGNED
------------------------------------------------------------------------------
The .docx carries no signature image. A signature sitting inside a Word file
can be pulled out in three clicks and pasted onto anything, and that file is
downloaded routinely by HR-OPS as part of ordinary work.

The signature is applied only when HR-APP approves the corrected PDF that comes
back. So the loop stays: generated PDF is signed -> Word copy is not -> the
corrected PDF HR-OPS uploads is not -> HR-APP's approval is what signs it.

------------------------------------------------------------------------------
THE CONTEXT DICTIONARY
------------------------------------------------------------------------------
    application_code      'DMRC-2026W-001'
    issued_on             date the letter was signed  -> "Dated:"
    salutation            'Mr.' | 'Ms.' | ''
    candidate_name        as stored, upper case
    course                degree programme, as stored
    college               college name, as stored
    duration_weeks        4 | 6 | 8
    sub_department        allotted posting, printed EXACTLY as stored
    start_date            actual date of joining
    end_date              last day, inclusive (see formatting.completion_date)
    session_term          'Winter' | 'Summer'
    application_year      2026
    signatory_name        HR-APP's name, as stored
    signatory_designation HR-APP's designation, as stored
    photo_path            passport photograph, or None
    signature_path        the signature image FROZEN at signing, or None

Any value may be missing. A letter with a blank in it is a visible problem that
HR-OPS will catch on screen; a letter that failed to generate at all is an
invisible one. So this module never raises on absent data -- the gate that
decides whether a letter MAY be issued lives on the server, not here.
"""

import io
import os

from reportlab.lib.colors import black
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas as pdf_canvas
from reportlab.platypus import Paragraph

from .formatting import (
    as_stored, format_date, duration_phrase, reference_line, subject_line,
)


# ==============================================================================
# PAGE GEOMETRY  (millimetres, measured from the format DMRC supplied)
#
# Adjust these after the first test print on real stationery. They are the only
# numbers that should ever need touching.
# ==============================================================================

PAGE_WIDTH, PAGE_HEIGHT = A4

# ------------------------------------------------------------------------------
# LETTERHEAD
#
# The portal draws the WHOLE letter, header included, so it prints on plain A4.
# No pre-printed stationery is required.
#
# Geometry measured from the format DMRC supplied. The logo sits at the left,
# the four title lines are centred on the space beside it -- which is why they
# centre on 95mm rather than the page's own centre of 105mm.
# ------------------------------------------------------------------------------

#: Folder holding the logo and the two pre-rendered Hindi lines.
ASSETS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'assets')

LOGO_FILENAME = 'dmrc_logo.png'
HINDI_LINE_1 = 'header_hindi_line1.png'
HINDI_LINE_2 = 'header_hindi_line2.png'

LOGO_LEFT_MM = 23.0
LOGO_TOP_MM = 15.0
LOGO_SIZE_MM = 20.0

#: Horizontal centre of the four title lines.
HEADER_CENTRE_MM = 95.0
HEADER_TOP_MM = 14.0

#: Height of each pre-rendered Hindi line on the page. Width follows from the
#: image's own proportions, so these never distort.
HINDI_1_HEIGHT_MM = 4.6
HINDI_2_HEIGHT_MM = 4.4

#: The rule beneath the letterhead.
HEADER_RULE_TOP_MM = 42.0

PAGE_WIDTH_MM = PAGE_WIDTH / mm

#: Where the body starts, below the letterhead rule at 42mm. THE VALUE TO
#: CHANGE if the body sits too close to, or too far below, that rule.
CONTENT_TOP_MM = 50.5

MARGIN_LEFT_MM = 22.0
MARGIN_RIGHT_MM = 22.0
MARGIN_BOTTOM_MM = 18.0

#: Word draws its own letterhead, so this is an ordinary top margin.
DOCX_TOP_MARGIN_MM = 14.0

#: Passport photograph box, top right. Measured off the supplied format.
PHOTO_BOX_LEFT_MM = 140.0
PHOTO_BOX_TOP_MM = 57.0
PHOTO_BOX_WIDTH_MM = 26.5
PHOTO_BOX_HEIGHT_MM = 29.7

#: Signature image, drawn above the signatory's name on the right.
SIGNATURE_WIDTH_MM = 38.0
SIGNATURE_HEIGHT_MM = 14.0

BODY_FONT = 'Helvetica'
BODY_FONT_BOLD = 'Helvetica-Bold'
BODY_SIZE = 10.5
LINE_LEADING = 15.0


def _x(mm_value):
    return mm_value * mm


def _y(mm_from_top):
    """Convert a distance from the TOP of the page into a PDF coordinate.

    PDF measures upward from the bottom left. Every measurement in this file is
    taken from the top, the way you would measure a sheet of paper, so this
    conversion happens once here instead of being scattered through the code.
    """
    return PAGE_HEIGHT - (mm_from_top * mm)


# ==============================================================================
# LETTERHEAD
# ==============================================================================

def _asset(filename):
    """Absolute path to a bundled asset, or None if it is not there.

    Returning None rather than raising means a missing logo produces a letter
    without a logo, not a failed generation and an intern left waiting.
    """
    path = os.path.join(ASSETS_DIR, filename)
    return path if os.path.exists(path) else None


def _draw_image_centred(canvas, image_path, centre_mm, top_mm, height_mm):
    """Draw an image at a fixed HEIGHT, centred, width following its own shape.

    Used for the pre-rendered Hindi lines. Setting only the height is what keeps
    them undistorted regardless of how the source images were produced.

    Returns the height actually drawn, or 0 if the file is unusable.
    """
    if not image_path:
        return 0.0
    try:
        from PIL import Image
        with Image.open(image_path) as img:
            src_w, src_h = img.size
    except Exception:
        return 0.0
    if not src_w or not src_h:
        return 0.0

    width_mm = height_mm * (src_w / src_h)
    try:
        canvas.drawImage(
            image_path,
            _x(centre_mm - width_mm / 2), _y(top_mm + height_mm),
            width=_x(width_mm), height=_x(height_mm),
            preserveAspectRatio=True, anchor='c', mask='auto',
        )
        return height_mm
    except Exception:
        return 0.0


def _draw_letterhead(canvas):
    """Draw the DMRC letterhead: logo, four title lines, and the rule beneath.

    WHY THE HINDI IS AN IMAGE
    -------------------------
    ReportLab cannot shape Devanagari. It places glyphs in the order the
    characters are stored, but Devanagari needs them reordered and joined --
    the vowel sign in "दिल्ली" is written before the consonant it follows, and
    "ल्ल" is a single conjunct form. Drawn as text, the words come out visibly
    wrong, and nobody reviewing an English-language codebase would spot it.

    So both Hindi lines are shipped as pre-rendered images in assets/. They are
    fixed text that will never change, they print at well over 1200 DPI, and
    they need no font file, no shaping library and no configuration on DMRC's
    server. The English lines are drawn as ordinary text, because Latin script
    needs none of that.

    TO CHANGE THE HINDI: the images must be re-rendered with a Devanagari font
    and proper shaping (Pillow with raqm does this correctly). Editing the
    strings in this file will not help -- they are not drawn from here.
    """
    logo = _asset(LOGO_FILENAME)
    if logo:
        _draw_image_fitted(
            canvas, logo,
            LOGO_LEFT_MM, LOGO_TOP_MM, LOGO_SIZE_MM, LOGO_SIZE_MM,
            draw_border=False,
        )

    y = HEADER_TOP_MM

    # Line 1 -- Hindi: (दिल्ली मेट्रो रेल कॉरपोरेशन लिमिटेड)
    drawn = _draw_image_centred(canvas, _asset(HINDI_LINE_1),
                                HEADER_CENTRE_MM, y, HINDI_1_HEIGHT_MM)
    y += (drawn or HINDI_1_HEIGHT_MM) + 1.4

    # Line 2 -- English name
    canvas.setFont('Helvetica-Bold', 12.5)
    canvas.drawCentredString(_x(HEADER_CENTRE_MM), _y(y + 4.2),
                             'DELHI METRO RAIL CORPORATION LIMITED')
    y += 6.0

    # Line 3 -- Hindi: (भारत सरकार एवं दिल्ली सरकार का संयुक्त उपक्रम)
    drawn = _draw_image_centred(canvas, _asset(HINDI_LINE_2),
                                HEADER_CENTRE_MM, y, HINDI_2_HEIGHT_MM)
    y += (drawn or HINDI_2_HEIGHT_MM) + 1.2

    # Line 4 -- the joint venture note, serif on the supplied format
    canvas.setFont('Times-Bold', 10.5)
    canvas.drawCentredString(_x(HEADER_CENTRE_MM), _y(y + 3.2),
                             '(A Joint Venture of Govt. of India and Govt. of Delhi)')
    y += 4.6

    # Line 5 -- the address
    canvas.setFont('Helvetica-Bold', 11)
    canvas.drawCentredString(_x(HEADER_CENTRE_MM), _y(y + 3.4),
                             'Metro Bhawan, New Delhi-110001')

    # The rule. Spans the body's own margins so the header and the text below
    # share one edge.
    canvas.setLineWidth(1.0)
    canvas.setStrokeColor(black)
    canvas.line(_x(MARGIN_LEFT_MM), _y(HEADER_RULE_TOP_MM),
                _x(PAGE_WIDTH_MM - MARGIN_RIGHT_MM), _y(HEADER_RULE_TOP_MM))


# ==============================================================================
# TEXT BLOCKS
#
# Written once, used by BOTH the PDF and the Word file, so the two cannot drift
# apart when the wording is revised. Each returns plain text plus the segments
# that must appear in bold.
# ==============================================================================

def _body_paragraph_1(ctx):
    salutation = as_stored(ctx.get('salutation'))
    name = as_stored(ctx.get('candidate_name'))
    who = f"{salutation} {name}".strip()
    course = as_stored(ctx.get('course'))
    college = as_stored(ctx.get('college'))
    period = duration_phrase(ctx.get('duration_weeks'))

    return (
        "With reference to the above note, approval has been granted for ",
        [
            (who, True),
            (", a ", False),
            (course, True),
            (" student at ", False),
            (college, True),
            (", to undergo training at DMRC for a period of ", False),
            (period, True),
            (". The Head of Department is requested to assign a live project "
             "to the trainee. ", False),
            ("The trainee is required to prepare a Project Report based on the "
             "live project assigned by the Head of Department", True),
            (".", False),
        ],
    )


def _body_paragraph_2(ctx):
    sub_dept = as_stored(ctx.get('sub_department'))
    start = format_date(ctx.get('start_date'))
    return [
        ("The trainee will be undergoing an internship in the ", False),
        (sub_dept, True),
        (" from ", False),
        (start, True),
        (".", False),
    ]


def _body_paragraph_3():
    return [
        ("The Head/Deputy Head of the department is requested to grant the "
         "trainee ", False),
        ("Limited Access", True),
        (" - restricted solely to relevant tasks or live projects.", False),
    ]


DISCLAIMER_TEXT = (
    "I understand that this internship is granted solely for academic purposes "
    "and is subject to my adherence to DMRC's policies, code of conduct, "
    "confidentiality requirements, and the terms and conditions of the "
    "internship."
)


def _segments_to_html(segments):
    """Turn (text, is_bold) pairs into the markup reportlab's Paragraph reads."""
    out = []
    for text, bold in segments:
        safe = (text or '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        out.append(f"<b>{safe}</b>" if bold else safe)
    return ''.join(out)


# ==============================================================================
# PDF
# ==============================================================================

def _draw_image_fitted(canvas, image_path, left_mm, top_mm, box_w_mm, box_h_mm,
                       draw_border=False):
    """Draw an image INSIDE a box without distorting it.

    Fitted, never cropped: the image is scaled until it touches the box on its
    longer side and centred on the other. A tall photograph gets white margins
    left and right rather than having somebody's head cut off, which is what
    filling the box would do.

    Returns False and draws nothing but the optional border if the file is
    missing or unreadable -- a cycle configured without a passport photograph
    leaves an empty box, exactly as agreed.
    """
    if draw_border:
        canvas.setLineWidth(0.8)
        canvas.setStrokeColor(black)
        canvas.rect(
            _x(left_mm), _y(top_mm + box_h_mm),
            _x(box_w_mm), _x(box_h_mm), stroke=1, fill=0,
        )

    if not image_path:
        return False

    try:
        from PIL import Image
        with Image.open(image_path) as img:
            src_w, src_h = img.size
    except Exception:
        return False

    if not src_w or not src_h:
        return False

    scale = min(box_w_mm / src_w, box_h_mm / src_h)
    draw_w, draw_h = src_w * scale, src_h * scale
    offset_x = left_mm + (box_w_mm - draw_w) / 2
    offset_y = top_mm + (box_h_mm - draw_h) / 2

    try:
        canvas.drawImage(
            image_path,
            _x(offset_x), _y(offset_y + draw_h),
            width=_x(draw_w), height=_x(draw_h),
            preserveAspectRatio=True, anchor='c', mask='auto',
        )
        return True
    except Exception:
        return False


def _draw_paragraph(canvas, html, top_mm, width_mm, style):
    """Lay out one wrapped paragraph and report how far down the page it ended."""
    para = Paragraph(html, style)
    _, height = para.wrap(_x(width_mm), _x(400))
    para.drawOn(canvas, _x(MARGIN_LEFT_MM), _y(top_mm) - height)
    return top_mm + (height / mm)


def build_offer_letter_pdf(ctx):
    """Return the offer letter as PDF bytes."""
    buffer = io.BytesIO()
    canvas = pdf_canvas.Canvas(buffer, pagesize=A4)
    canvas.setTitle(f"Offer Letter {as_stored(ctx.get('application_code'))}")
    canvas.setAuthor('Delhi Metro Rail Corporation Limited')

    content_width_mm = PAGE_WIDTH / mm - MARGIN_LEFT_MM - MARGIN_RIGHT_MM

    body = ParagraphStyle(
        'body', fontName=BODY_FONT, fontSize=BODY_SIZE, leading=LINE_LEADING,
        alignment=TA_JUSTIFY, textColor=black,
    )
    centred = ParagraphStyle('centred', parent=body, alignment=TA_CENTER)
    right = ParagraphStyle('right', parent=body, alignment=TA_RIGHT)

    _draw_letterhead(canvas)

    y = CONTENT_TOP_MM

    # --- Reference line and date, on one line at opposite margins ------------
    canvas.setFont(BODY_FONT_BOLD, BODY_SIZE)
    canvas.drawString(_x(MARGIN_LEFT_MM), _y(y),
                      reference_line(ctx.get('application_code')))
    canvas.drawRightString(_x(PAGE_WIDTH / mm - MARGIN_RIGHT_MM), _y(y),
                           f"Dated: {format_date(ctx.get('issued_on'))}")
    y += 11

    # --- "Note", centred and underlined --------------------------------------
    note_centre = MARGIN_LEFT_MM + content_width_mm / 2
    canvas.setFont(BODY_FONT_BOLD, BODY_SIZE)
    canvas.drawCentredString(_x(note_centre), _y(y), 'Note')
    note_half = canvas.stringWidth('Note', BODY_FONT_BOLD, BODY_SIZE) / 2
    canvas.setLineWidth(0.7)
    canvas.line(_x(note_centre) - note_half, _y(y) - 1.8,
                _x(note_centre) + note_half, _y(y) - 1.8)
    y += 11

    # --- Subject --------------------------------------------------------------
    subject = subject_line(ctx.get('duration_weeks'), ctx.get('session_term'),
                           ctx.get('application_year'))
    canvas.setFont(BODY_FONT_BOLD, BODY_SIZE)
    canvas.drawString(_x(MARGIN_LEFT_MM), _y(y), f"Subject: {subject}")

    # --- Passport photograph, top right --------------------------------------
    # Drawn at a fixed position rather than in the flow of the text, because it
    # sits beside the subject line on the supplied format.
    _draw_image_fitted(
        canvas, ctx.get('photo_path'),
        PHOTO_BOX_LEFT_MM, PHOTO_BOX_TOP_MM,
        PHOTO_BOX_WIDTH_MM, PHOTO_BOX_HEIGHT_MM,
        draw_border=True,
    )

    # Clear the photo box before resuming body text, whichever is lower.
    y = max(y + 12, PHOTO_BOX_TOP_MM + PHOTO_BOX_HEIGHT_MM + 8)

    # --- Paragraph 1 ----------------------------------------------------------
    lead, segments = _body_paragraph_1(ctx)
    html = lead + _segments_to_html(segments)
    y = _draw_paragraph(canvas, html, y, content_width_mm, body) + 7

    # --- Schedule table -------------------------------------------------------
    canvas.setFont(BODY_FONT, BODY_SIZE)
    canvas.drawString(_x(MARGIN_LEFT_MM), _y(y), 'The training schedule is outlined below:')
    y += 7

    col_widths = [18.0, 52.0, content_width_mm - 70.0]
    row_height = 9.0
    table_left = MARGIN_LEFT_MM

    period = ''
    if ctx.get('start_date') and ctx.get('end_date'):
        period = f"{format_date(ctx['start_date'])} – {format_date(ctx['end_date'])}"
    elif ctx.get('start_date'):
        period = format_date(ctx['start_date'])

    rows = [
        ('Sr. No.', 'Deputed Under', 'Period', True),
        ('1.', as_stored(ctx.get('sub_department')), period, False),
    ]

    for label_a, label_b, label_c, is_header in rows:
        canvas.setLineWidth(0.8)
        x_cursor = table_left
        for width, text in zip(col_widths, (label_a, label_b, label_c)):
            canvas.rect(_x(x_cursor), _y(y + row_height), _x(width), _x(row_height),
                        stroke=1, fill=0)
            canvas.setFont(BODY_FONT_BOLD if is_header else BODY_FONT, BODY_SIZE)
            canvas.drawCentredString(
                _x(x_cursor + width / 2), _y(y + row_height - 3.0), text or '',
            )
            x_cursor += width
        y += row_height
    y += 10

    # --- Paragraph 2 ----------------------------------------------------------
    y = _draw_paragraph(canvas, _segments_to_html(_body_paragraph_2(ctx)),
                        y, content_width_mm, body) + 8

    # --- Paragraph 3 ----------------------------------------------------------
    y = _draw_paragraph(canvas, _segments_to_html(_body_paragraph_3()),
                        y, content_width_mm, body) + 6

    # --- Signature block, right aligned --------------------------------------
    sig_right = PAGE_WIDTH / mm - MARGIN_RIGHT_MM
    sig_left = sig_right - SIGNATURE_WIDTH_MM

    if ctx.get('signature_path'):
        _draw_image_fitted(
            canvas, ctx['signature_path'],
            sig_left, y, SIGNATURE_WIDTH_MM, SIGNATURE_HEIGHT_MM,
            draw_border=False,
        )
    y += SIGNATURE_HEIGHT_MM + 1

    canvas.setLineWidth(0.7)
    canvas.line(_x(sig_left - 8), _y(y), _x(sig_right), _y(y))
    y += 6

    canvas.setFont(BODY_FONT, BODY_SIZE)
    canvas.drawRightString(_x(sig_right), _y(y), as_stored(ctx.get('signatory_name')))
    y += 6
    designation = as_stored(ctx.get('signatory_designation'))
    canvas.drawRightString(_x(sig_right), _y(y), f"{designation}/HR" if designation else 'HR')
    y += 14

    # --- Disclaimer -----------------------------------------------------------
    # First person, and signed by the INTERN. Deliberately left blank below:
    # they complete it by hand when they collect the letter.
    disclaimer_html = f"<b><u>Disclaimer</u></b>: {DISCLAIMER_TEXT}"
    y = _draw_paragraph(canvas, disclaimer_html, y, content_width_mm, body) + 12

    canvas.setFont(BODY_FONT_BOLD, BODY_SIZE)
    canvas.drawString(_x(MARGIN_LEFT_MM), _y(y), 'Name:')
    canvas.drawString(_x(MARGIN_LEFT_MM + 58), _y(y), 'Place:')
    canvas.drawString(_x(MARGIN_LEFT_MM + 106), _y(y), 'Date:')
    y += 10
    canvas.drawString(_x(MARGIN_LEFT_MM), _y(y), 'Signature:')
    sig_label_width = canvas.stringWidth('Signature: ', BODY_FONT_BOLD, BODY_SIZE)
    canvas.setLineWidth(0.7)
    canvas.line(_x(MARGIN_LEFT_MM) + sig_label_width, _y(y) - 1,
                _x(MARGIN_LEFT_MM + 88), _y(y) - 1)

    canvas.showPage()
    canvas.save()
    return buffer.getvalue()


# ==============================================================================
# WORD
#
# Same content, no signature. Produced independently of the PDF because there is
# no dependable way to convert between the two without installing extra software
# on DMRC's server -- which the handover deliberately avoids.
#
# The two therefore have to be kept in step BY HAND. The shared text blocks
# above are what makes that manageable: revise the wording there and both files
# change together. Anything altered directly in one of the two builders will
# drift.
# ==============================================================================

def _docx_letterhead(document):
    """Draw the letterhead into the Word file: logo left, titles beside it.

    Laid out as a borderless two-cell table because that is the only reliable
    way to put an image beside text in a .docx -- Word's floating images move
    unpredictably when the text around them is edited, and HR-OPS edits this
    file by design.

    The Hindi lines are the same pre-rendered images the PDF uses, for the same
    reason: see _draw_letterhead().
    """
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm, Pt

    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    logo_cell, title_cell = table.cell(0, 0), table.cell(0, 1)
    logo_cell.width = Mm(28)
    title_cell.width = Mm(138)

    logo = _asset(LOGO_FILENAME)
    if logo:
        para = logo_cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            para.add_run().add_picture(logo, height=Mm(LOGO_SIZE_MM))
        except Exception:
            pass

    def title_line(text=None, image=None, height_mm=None, size=None,
                   font=None, first=False):
        para = title_cell.paragraphs[0] if first else title_cell.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        if image:
            try:
                para.add_run().add_picture(image, height=Mm(height_mm))
                return
            except Exception:
                return
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        if font:
            run.font.name = font

    title_line(image=_asset(HINDI_LINE_1), height_mm=HINDI_1_HEIGHT_MM, first=True)
    title_line(text='DELHI METRO RAIL CORPORATION LIMITED', size=12.5)
    title_line(image=_asset(HINDI_LINE_2), height_mm=HINDI_2_HEIGHT_MM)
    title_line(text='(A Joint Venture of Govt. of India and Govt. of Delhi)',
               size=10.5, font='Times New Roman')
    title_line(text='Metro Bhawan, New Delhi-110001', size=11)

    _docx_horizontal_rule(document.add_paragraph())


def _docx_horizontal_rule(paragraph):
    """Give a paragraph a bottom border -- the rule beneath the letterhead.

    python-docx has no API for this, so the border element is added to the
    paragraph's XML directly. This is the documented way to do it.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    borders.append(bottom)
    properties.append(borders)


def _docx_runs(paragraph, segments):
    for text, bold in segments:
        if not text:
            continue
        run = paragraph.add_run(text)
        run.bold = bold


def build_offer_letter_docx(ctx):
    """Return the offer letter as .docx bytes, WITHOUT the signature."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm, Pt

    document = Document()

    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    # The Word file draws its own letterhead, exactly as the PDF does, so the
    # margin here is an ordinary one rather than a gap left for pre-printing.
    section.top_margin = Mm(DOCX_TOP_MARGIN_MM)
    section.bottom_margin = Mm(MARGIN_BOTTOM_MM)
    section.left_margin = Mm(MARGIN_LEFT_MM)
    section.right_margin = Mm(MARGIN_RIGHT_MM)

    normal = document.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(BODY_SIZE)

    _docx_letterhead(document)

    # --- Reference and date, tab-separated to opposite margins ---------------
    from docx.enum.text import WD_TAB_ALIGNMENT
    head = document.add_paragraph()
    head.paragraph_format.tab_stops.add_tab_stop(Mm(166), WD_TAB_ALIGNMENT.RIGHT)
    run = head.add_run(reference_line(ctx.get('application_code')))
    run.bold = True
    run = head.add_run(f"\tDated: {format_date(ctx.get('issued_on'))}")
    run.bold = True

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run('Note')
    run.bold = True
    run.underline = True

    subject = document.add_paragraph()
    run = subject.add_run(
        'Subject: ' + subject_line(ctx.get('duration_weeks'),
                                   ctx.get('session_term'),
                                   ctx.get('application_year'))
    )
    run.bold = True

    # The photograph is embedded so the corrected letter still carries it.
    # Unlike the signature, a passport photograph is the candidate's own and
    # is already held in the document vault -- extracting it from the Word file
    # gains nobody anything they could not already see.
    if ctx.get('photo_path'):
        try:
            photo_para = document.add_paragraph()
            photo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            photo_para.add_run().add_picture(
                ctx['photo_path'], height=Mm(PHOTO_BOX_HEIGHT_MM),
            )
        except Exception:
            pass

    document.add_paragraph()

    lead, segments = _body_paragraph_1(ctx)
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.add_run(lead)
    _docx_runs(para, segments)

    document.add_paragraph('The training schedule is outlined below:')

    period = ''
    if ctx.get('start_date') and ctx.get('end_date'):
        period = f"{format_date(ctx['start_date'])} – {format_date(ctx['end_date'])}"
    elif ctx.get('start_date'):
        period = format_date(ctx['start_date'])

    table = document.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    for index, heading in enumerate(('Sr. No.', 'Deputed Under', 'Period')):
        cell = table.cell(0, index)
        cell.text = ''
        run = cell.paragraphs[0].add_run(heading)
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, value in enumerate(('1.', as_stored(ctx.get('sub_department')), period)):
        cell = table.cell(1, index)
        cell.text = ''
        cell.paragraphs[0].add_run(value or '')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()

    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _docx_runs(para, _body_paragraph_2(ctx))

    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _docx_runs(para, _body_paragraph_3())

    document.add_paragraph()

    # --- Signature block: the LINE and the name, but never the image ---------
    for text in ('______________________',
                 as_stored(ctx.get('signatory_name')),
                 (f"{as_stored(ctx.get('signatory_designation'))}/HR"
                  if ctx.get('signatory_designation') else 'HR')):
        para = document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para.add_run(text)

    document.add_paragraph()

    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run('Disclaimer')
    run.bold = True
    run.underline = True
    para.add_run(': ' + DISCLAIMER_TEXT)

    document.add_paragraph()

    footer = document.add_paragraph()
    footer.paragraph_format.tab_stops.add_tab_stop(Mm(58), WD_TAB_ALIGNMENT.LEFT)
    footer.paragraph_format.tab_stops.add_tab_stop(Mm(106), WD_TAB_ALIGNMENT.LEFT)
    run = footer.add_run('Name:\tPlace:\tDate:')
    run.bold = True

    signature = document.add_paragraph()
    run = signature.add_run('Signature: _______________________________')
    run.bold = True

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
