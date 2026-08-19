"""
COMPLETION CERTIFICATE GENERATION
==============================================================================

    build_completion_certificate_pdf(ctx)    the signed certificate, for print
    build_completion_certificate_docx(ctx)   an editable copy for corrections

Both return raw bytes. Neither touches the database, the filesystem or Django
settings, and neither imports a model -- the same boundary the offer letter
keeps, and what lets the layout be checked by running
preview_completion_certificate.py with no database and no server.

The LETTERHEAD is shared with the offer letter, drawn by the same code from the
same measurements. Change it once and both documents change together.

------------------------------------------------------------------------------
WHY THE WORD COPY IS UNSIGNED
------------------------------------------------------------------------------
Identical reasoning to the offer letter. HR-APP downloads the Word copy to
correct a certificate, and a signature sitting inside a downloadable Word file
can be lifted in three clicks and pasted onto anything. The corrected PDF gets
its signature when HR-APP approves it, applied by documents/stamp.py.

------------------------------------------------------------------------------
THE CONTEXT DICTIONARY
------------------------------------------------------------------------------
    application_code      'DMRC-2026W-001'  -> the reference line, same as the
                          offer letter, so both documents for one candidate
                          carry the same number
    issued_on             date the certificate was signed -> "Dated:"
    salutation            'Mr.' | 'Ms.' | ''
    candidate_name        as stored, upper case
    college               college name, as stored
    sub_department        the posting, printed EXACTLY as stored
    start_date            actual date of joining
    end_date              date_of_completion -- the SAME date the offer letter
                          printed, not a fresh calculation
    project_title         printed in quotation marks
    gender                chooses the pronouns; see formatting.pronouns_for()
    signatory_name        HR-APP's name, as stored
    signatory_designation HR-APP's designation, as stored
    signature_path        the signature FROZEN at signing, or None

NOT on the certificate, deliberately: the file number. It is recorded against
the application and shown in the drawer, and DMRC does not print it here.

Any value may be missing. A certificate with a blank in it is a visible problem
that HR-APP will catch on screen; one that failed to generate is an invisible
one. The gate deciding whether a certificate MAY be issued lives on the server.
"""

import io

from reportlab.lib.colors import black
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas as pdf_canvas

from .formatting import as_stored, format_date, pronouns_for, reference_line
# The letterhead, the page geometry and the drawing helpers are the offer
# letter's. Shared rather than copied, so the two documents cannot drift apart.
from .offer_letter import (
    BODY_FONT, BODY_FONT_BOLD, BODY_SIZE, CONTENT_TOP_MM, DOCX_TOP_MARGIN_MM,
    LINE_LEADING, MARGIN_BOTTOM_MM, MARGIN_LEFT_MM, MARGIN_RIGHT_MM,
    PAGE_HEIGHT, PAGE_WIDTH, PAGE_WIDTH_MM, SIGNATURE_HEIGHT_MM,
    SIGNATURE_WIDTH_MM, _docx_letterhead, _draw_image_fitted, _draw_letterhead,
    _draw_paragraph, _segments_to_html, _x, _y,
)


# ==============================================================================
# TEXT BLOCKS
#
# Written once and used by BOTH builders, so the PDF and the Word copy cannot
# drift when the wording is revised. Each returns (text, is_bold) segments.
# ==============================================================================

HEADING = 'TO WHOMSOEVER IT MAY CONCERN'

ACADEMIC_PURPOSE_NOTE = '(This certificate is being issued for Academic purpose only)'


def _paragraph_1(ctx):
    """Who, where, and for how long."""
    salutation = as_stored(ctx.get('salutation'))
    name = as_stored(ctx.get('candidate_name'))
    who = f"{salutation} {name}".strip()
    p = pronouns_for(ctx.get('gender'))

    period = ''
    if ctx.get('start_date') and ctx.get('end_date'):
        period = f"{format_date(ctx['start_date'])} to {format_date(ctx['end_date'])}"

    return [
        ('This is to certify that ', False),
        (who, True),
        (', student of ', False),
        (as_stored(ctx.get('college')), True),
        (f" has completed {p['possessive']} internship in ", False),
        (as_stored(ctx.get('sub_department')), True),
        (' of ', False),
        ('DMRC', True),
        (' for a period from ', False),
        (format_date(ctx.get('start_date')), True),
        (' to ', False),
        (format_date(ctx.get('end_date')), True),
        (' as a part of educational course curriculum.', False),
    ] if period else [
        ('This is to certify that ', False),
        (who, True),
        (', student of ', False),
        (as_stored(ctx.get('college')), True),
        (f" has completed {p['possessive']} internship in ", False),
        (as_stored(ctx.get('sub_department')), True),
        (' of ', False),
        ('DMRC', True),
        (' as a part of educational course curriculum.', False),
    ]


def _paragraph_2(ctx):
    """The project. The title is printed inside quotation marks."""
    p = pronouns_for(ctx.get('gender'))
    title = as_stored(ctx.get('project_title'))
    return [
        (f"During the internship, {p['subject']} successfully worked on the "
         f"project titled ", False),
        (f'"{title}"', True),
        (' as part of the assigned departmental objectives and educational '
         'course curriculum.', False),
    ]


def _paragraph_3(ctx):
    """The commendation.

    Fixed text with no room for a mentor's actual assessment -- which is
    precisely why an Unsatisfactory evaluation produces no certificate at all.
    There is no honest version of this paragraph for a failed internship.
    """
    p = pronouns_for(ctx.get('gender'))
    return [
        (f"{p['possessive_cap']} approach towards the training was enthusiastic. "
         f"During this period, {p['subject']} was found to be sincere, "
         f"pro-active and hard-working.", False),
    ]


def _paragraph_4(ctx):
    p = pronouns_for(ctx.get('gender'))
    return [
        (f"Delhi Metro Rail Corporation wishes {p['object']} success in all "
         f"{p['possessive']} future endeavors.", False),
    ]


# ==============================================================================
# PDF
# ==============================================================================

def build_completion_certificate_pdf(ctx):
    """Return the completion certificate as PDF bytes."""
    buffer = io.BytesIO()
    canvas = pdf_canvas.Canvas(buffer, pagesize=A4)
    canvas.setTitle(f"Completion Certificate {as_stored(ctx.get('application_code'))}")
    canvas.setAuthor('Delhi Metro Rail Corporation Limited')

    content_width_mm = PAGE_WIDTH_MM - MARGIN_LEFT_MM - MARGIN_RIGHT_MM

    body = ParagraphStyle(
        'body', fontName=BODY_FONT, fontSize=BODY_SIZE, leading=LINE_LEADING,
        alignment=TA_JUSTIFY, textColor=black,
    )
    centred = ParagraphStyle('centred', parent=body, alignment=TA_CENTER)

    _draw_letterhead(canvas)

    y = CONTENT_TOP_MM

    # --- Reference and date, opposite margins --------------------------------
    canvas.setFont(BODY_FONT_BOLD, BODY_SIZE)
    canvas.drawString(_x(MARGIN_LEFT_MM), _y(y),
                      reference_line(ctx.get('application_code')))
    canvas.drawRightString(_x(PAGE_WIDTH_MM - MARGIN_RIGHT_MM), _y(y),
                           f"Dated: {format_date(ctx.get('issued_on'))}")
    y += 16

    # --- Heading, centred and underlined -------------------------------------
    centre = MARGIN_LEFT_MM + content_width_mm / 2
    canvas.setFont(BODY_FONT_BOLD, BODY_SIZE)
    canvas.drawCentredString(_x(centre), _y(y), HEADING)
    half = canvas.stringWidth(HEADING, BODY_FONT_BOLD, BODY_SIZE) / 2
    canvas.setLineWidth(0.7)
    canvas.line(_x(centre) - half, _y(y) - 1.8, _x(centre) + half, _y(y) - 1.8)
    y += 12

    # --- The four paragraphs --------------------------------------------------
    for segments in (_paragraph_1(ctx), _paragraph_2(ctx),
                     _paragraph_3(ctx), _paragraph_4(ctx)):
        y = _draw_paragraph(canvas, _segments_to_html(segments),
                            y, content_width_mm, body) + 7

    y += 4

    # --- The academic-purpose note, centred and bold --------------------------
    canvas.setFont(BODY_FONT_BOLD, BODY_SIZE)
    canvas.drawCentredString(_x(centre), _y(y), ACADEMIC_PURPOSE_NOTE)
    y += 26

    # --- Signature block, right aligned ---------------------------------------
    sig_right = PAGE_WIDTH_MM - MARGIN_RIGHT_MM
    sig_left = sig_right - SIGNATURE_WIDTH_MM

    if ctx.get('signature_path'):
        _draw_image_fitted(
            canvas, ctx['signature_path'],
            sig_left, y, SIGNATURE_WIDTH_MM, SIGNATURE_HEIGHT_MM,
            draw_border=False,
        )
    y += SIGNATURE_HEIGHT_MM + 1

    canvas.setLineWidth(0.7)
    canvas.line(_x(sig_left - 8), _y(y), _x(sig_right), _y(y))
    y += 6

    canvas.setFont(BODY_FONT, BODY_SIZE)
    canvas.drawRightString(_x(sig_right), _y(y), as_stored(ctx.get('signatory_name')))
    y += 6
    designation = as_stored(ctx.get('signatory_designation'))
    canvas.drawRightString(_x(sig_right), _y(y),
                           f"{designation}/HR" if designation else 'HR')

    canvas.showPage()
    canvas.save()
    return buffer.getvalue()


# ==============================================================================
# WORD
#
# Same content, NO signature. Built independently of the PDF because there is no
# dependable way to convert between the two without installing extra software on
# DMRC's server, which the handover deliberately avoids.
#
# The shared text blocks above are what keeps the two in step: revise the
# wording there and both change together. Anything edited directly in one
# builder will drift.
# ==============================================================================

def _docx_runs(paragraph, segments):
    for text, bold in segments:
        if not text:
            continue
        run = paragraph.add_run(text)
        run.bold = bold


def build_completion_certificate_docx(ctx):
    """Return the certificate as .docx bytes, WITHOUT the signature."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.shared import Mm, Pt

    document = Document()

    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(DOCX_TOP_MARGIN_MM)
    section.bottom_margin = Mm(MARGIN_BOTTOM_MM)
    section.left_margin = Mm(MARGIN_LEFT_MM)
    section.right_margin = Mm(MARGIN_RIGHT_MM)

    normal = document.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(BODY_SIZE)

    _docx_letterhead(document)

    head = document.add_paragraph()
    head.paragraph_format.tab_stops.add_tab_stop(Mm(166), WD_TAB_ALIGNMENT.RIGHT)
    run = head.add_run(reference_line(ctx.get('application_code')))
    run.bold = True
    run = head.add_run(f"\t Dated: {format_date(ctx.get('issued_on'))}")
    run.bold = True

    document.add_paragraph()

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(HEADING)
    run.bold = True
    run.underline = True

    document.add_paragraph()

    for segments in (_paragraph_1(ctx), _paragraph_2(ctx),
                     _paragraph_3(ctx), _paragraph_4(ctx)):
        para = document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _docx_runs(para, segments)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(ACADEMIC_PURPOSE_NOTE)
    run.bold = True

    for _ in range(3):
        document.add_paragraph()

    # The signature LINE and the name, but never the image. The corrected PDF is
    # signed when HR-APP approves it, not before.
    for text in ('______________________',
                 as_stored(ctx.get('signatory_name')),
                 (f"{as_stored(ctx.get('signatory_designation'))}/HR"
                  if ctx.get('signatory_designation') else 'HR')):
        para = document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para.add_run(text)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
